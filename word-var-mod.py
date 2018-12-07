from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

document = Document()

### Entrada de Datos para Variables
nombreDelCliente = input("Nombre del Cliente: ")
rucDelCliente = input("Ruc del Cliente: ")
direccionDelCliente = input("Direccion del Cliente: ")
apoderadoDelCliente = input("Apoderado del Cliente: ")
documentoDelCliente = input("Documento del Cliente: ")
###

############### Style Arial 11
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
###############

### Titulo
titulo = document.add_paragraph(" ")
titulo.add_run("CONTRATO DE PRESTACIÓN DE SERVICIOS DE CONSULTORÍA COMERCIAL EN CANALES DIGITALES\n").bold = True
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER


### Primer Parrafo
parrafo1 = document.add_paragraph("Conste por el presente documento el Contrato de Prestación de Servicios de Consultoría Comercial en Canales Digitales (el “Contrato”), que celebran de una parte") 
parrafo1.add_run(f" {nombreDelCliente} ").bold = True
parrafo1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
parrafo1.add_run("identificado con RUC N° ")
parrafo1.add_run(f"{rucDelCliente}").bold = True
parrafo1.add_run(f", con domicilio en ")
parrafo1.add_run(f"{direccionDelCliente} ").bold = True
parrafo1.add_run(", debidamente representada por sus Apoderados, señor/a ")
parrafo1.add_run(f"{apoderadoDelCliente} ").bold = True
parrafo1.add_run("con Cédula de Identidad N° ")
parrafo1.add_run(f"{documentoDelCliente}").bold = True
parrafo1.add_run(", a quien en adelante se le denominará ")
parrafo1.add_run("EL CLIENTE").bold = True
parrafo1.add_run("; y, de la otra parte ")
parrafo1.add_run("MESTIZA.").bold = True
parrafo1.add_run(", identificada con RUC N° 20602639160, con domicilio en Julián Arias Araguez 425, dpto 1003, Provincia y Departamento de Lima, debidamente ")
parrafo1.add_run("representada por su Gerente General, señor Armando Fernández Mendieta, identificado con DNI N° 47101934, ")
parrafo1.add_run("conforme consta en poder inscrito en la Partida N° 13979407, a quien en adelante se le denominará ")
parrafo1.add_run("MESTIZA (EL CLIENTE y MESTIZA,").bold = True
parrafo1.add_run("de manera conjunta, se denominarán las “Partes”), en los términos y condiciones siguientes: \n")

### Segundo Parrafo

clausula1 = document.add_paragraph(" ")
clausula1.add_run("CLÁUSULA PRIMERA:\t ANTECEDENTES").bold = True





document.add_page_break()

document.save('Contrato/Contrato.docx')